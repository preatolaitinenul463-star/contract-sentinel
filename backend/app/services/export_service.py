"""Report export service - generates Word/PDF reports and redline annotations."""
import io
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime

from loguru import logger

from app.config import settings
from app.models.review import ReviewResult
from app.models.contract import Contract


class ExportService:
    """Service for exporting review reports."""
    
    async def export_review_docx(
        self,
        contract: Contract,
        review: ReviewResult,
        output_dir: Path = None,
    ) -> str:
        """Export review result as Word document."""
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        if output_dir is None:
            output_dir = settings.storage_dir / "reports"
        output_dir.mkdir(parents=True, exist_ok=True)
        
        doc = Document()
        
        # Title
        title = doc.add_heading("合同审核报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Basic info
        doc.add_heading("基本信息", level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = "Table Grid"
        
        info_data = [
            ("合同名称", contract.filename),
            ("合同类型", contract.contract_type.value),
            ("适用法域", contract.jurisdiction.value),
            ("审核时间", review.created_at.strftime("%Y-%m-%d %H:%M:%S")),
        ]
        
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = str(value)
        
        # Summary
        doc.add_heading("风险概览", level=1)
        
        summary_para = doc.add_paragraph()
        summary_para.add_run(f"高风险: {review.high_risk_count}  ").bold = True
        summary_para.add_run(f"中风险: {review.medium_risk_count}  ")
        summary_para.add_run(f"低风险: {review.low_risk_count}")
        
        if review.summary:
            doc.add_paragraph(review.summary)
        
        # Risk items
        doc.add_heading("风险条款详情", level=1)
        
        severity_colors = {
            "high": RGBColor(220, 53, 69),
            "medium": RGBColor(255, 193, 7),
            "low": RGBColor(23, 162, 184),
        }
        
        for i, risk in enumerate(review.risk_items or [], 1):
            # Risk title
            risk_heading = doc.add_heading(f"{i}. {risk.get('name', '未知风险')}", level=2)
            
            # Severity
            severity = risk.get("severity", "medium")
            severity_para = doc.add_paragraph()
            severity_run = severity_para.add_run(f"[{severity.upper()}] ")
            severity_run.font.color.rgb = severity_colors.get(severity, RGBColor(0, 0, 0))
            severity_run.bold = True
            severity_para.add_run(risk.get("description", ""))
            
            # Original clause
            if risk.get("clause_text"):
                doc.add_paragraph("原文条款：", style="Intense Quote")
                clause_para = doc.add_paragraph(risk["clause_text"])
                clause_para.paragraph_format.left_indent = Inches(0.5)
            
            # Suggestion
            if risk.get("suggestion"):
                doc.add_paragraph("修改建议：", style="Intense Quote")
                suggestion_para = doc.add_paragraph(risk["suggestion"])
                suggestion_para.paragraph_format.left_indent = Inches(0.5)
            
            doc.add_paragraph()  # Spacer
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.add_run("本报告由合同哨兵AI生成，仅供参考，不构成法律意见或律师执业服务。重要决策请咨询持证律师。").italic = True
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save
        filename = f"review_{contract.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = output_dir / filename
        doc.save(output_path)
        
        return str(output_path)
    
    async def export_review_pdf(
        self,
        contract: Contract,
        review: ReviewResult,
        output_dir: Path = None,
    ) -> str:
        """Export review result as PDF document."""
        # First generate DOCX, then convert to PDF
        # For simplicity, we'll generate HTML and convert with WeasyPrint
        
        if output_dir is None:
            output_dir = settings.storage_dir / "reports"
        output_dir.mkdir(parents=True, exist_ok=True)
        
        html_content = self._generate_html_report(contract, review)
        
        try:
            from weasyprint import HTML
            
            filename = f"review_{contract.id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            output_path = output_dir / filename
            
            HTML(string=html_content).write_pdf(output_path)
            
            return str(output_path)
        except ImportError:
            logger.warning("WeasyPrint not available, falling back to DOCX")
            return await self.export_review_docx(contract, review, output_dir)
    
    def _generate_html_report(self, contract: Contract, review: ReviewResult) -> str:
        """Generate HTML report content."""
        severity_colors = {
            "high": "#dc3545",
            "medium": "#ffc107",
            "low": "#17a2b8",
        }
        
        risk_html = ""
        for i, risk in enumerate(review.risk_items or [], 1):
            severity = risk.get("severity", "medium")
            color = severity_colors.get(severity, "#666")
            
            risk_html += f"""
            <div class="risk-item">
                <h3>{i}. {risk.get('name', '未知风险')}</h3>
                <p><span class="severity" style="background-color: {color};">{severity.upper()}</span>
                   {risk.get('description', '')}</p>
                {'<div class="clause"><strong>原文条款：</strong><p>' + risk['clause_text'] + '</p></div>' if risk.get('clause_text') else ''}
                {'<div class="suggestion"><strong>修改建议：</strong><p>' + risk['suggestion'] + '</p></div>' if risk.get('suggestion') else ''}
            </div>
            """
        
        html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>合同审核报告</title>
            <style>
                body {{ font-family: "Microsoft YaHei", sans-serif; padding: 20px; }}
                h1 {{ text-align: center; color: #333; }}
                .info-table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
                .info-table td {{ border: 1px solid #ddd; padding: 8px; }}
                .info-table td:first-child {{ background: #f5f5f5; width: 120px; }}
                .summary {{ background: #f8f9fa; padding: 15px; border-radius: 5px; }}
                .risk-item {{ margin: 20px 0; padding: 15px; border: 1px solid #ddd; border-radius: 5px; }}
                .severity {{ color: white; padding: 2px 8px; border-radius: 3px; font-size: 12px; }}
                .clause {{ background: #fff3cd; padding: 10px; margin: 10px 0; border-radius: 3px; }}
                .suggestion {{ background: #d4edda; padding: 10px; margin: 10px 0; border-radius: 3px; }}
                .footer {{ text-align: center; color: #666; font-style: italic; margin-top: 30px; }}
            </style>
        </head>
        <body>
            <h1>合同审核报告</h1>
            
            <h2>基本信息</h2>
            <table class="info-table">
                <tr><td>合同名称</td><td>{contract.filename}</td></tr>
                <tr><td>合同类型</td><td>{contract.contract_type.value}</td></tr>
                <tr><td>适用法域</td><td>{contract.jurisdiction.value}</td></tr>
                <tr><td>审核时间</td><td>{review.created_at.strftime('%Y-%m-%d %H:%M:%S')}</td></tr>
            </table>
            
            <h2>风险概览</h2>
            <div class="summary">
                <p><strong>高风险:</strong> {review.high_risk_count} &nbsp;
                   <strong>中风险:</strong> {review.medium_risk_count} &nbsp;
                   <strong>低风险:</strong> {review.low_risk_count}</p>
                <p>{review.summary or ''}</p>
            </div>
            
            <h2>风险条款详情</h2>
            {risk_html}
            
            <p class="footer">本报告由合同哨兵AI生成，仅供参考，不构成法律意见或律师执业服务。重要决策请咨询持证律师。</p>
        </body>
        </html>
        """
        
        return html

    # ═══════════════════════════════════════════════════════════
    # 法律助手报告导出
    # ═══════════════════════════════════════════════════════════

    async def export_assistant_report_docx(
        self,
        run,
        sources: list,
        mode_label: str = "法律分析",
        output_dir: Path = None,
    ) -> str:
        """Export assistant report as Word document with sources appendix."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        if output_dir is None:
            output_dir = settings.storage_dir / "reports"
        output_dir.mkdir(parents=True, exist_ok=True)

        doc = Document()

        # Title
        title = doc.add_heading(f"法律助手 · {mode_label}报告", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Info
        doc.add_heading("基本信息", level=1)
        info_table = doc.add_table(rows=4, cols=2)
        info_table.style = "Table Grid"
        info_data = [
            ("报告类型", mode_label),
            ("运行编号", run.run_id if hasattr(run, 'run_id') else str(run.id)),
            ("法域", run.jurisdiction or "CN"),
            ("生成时间", run.created_at.strftime("%Y-%m-%d %H:%M:%S") if run.created_at else ""),
        ]
        for i, (label, value) in enumerate(info_data):
            info_table.rows[i].cells[0].text = label
            info_table.rows[i].cells[1].text = str(value)

        # Summary
        if run.result_summary:
            doc.add_heading("结果概览", level=1)
            summary = run.result_summary
            if isinstance(summary, dict):
                for k, v in summary.items():
                    doc.add_paragraph(f"{k}: {v}")

        # Sources appendix
        if sources:
            doc.add_heading("来源附录", level=1)
            for s in sources:
                sid = s.get("source_id", "")
                trusted_label = "【官方】" if s.get("trusted") else "【参考】"
                source_title = s.get("title", "无标题")

                sp = doc.add_paragraph()
                sr = sp.add_run(f"[{sid}] {trusted_label} {source_title}")
                sr.bold = True
                sr.font.size = Pt(10)
                if s.get("trusted"):
                    sr.font.color.rgb = RGBColor(0, 128, 0)
                else:
                    sr.font.color.rgb = RGBColor(200, 150, 0)

                if s.get("url"):
                    doc.add_paragraph(f"链接: {s['url']}")
                if s.get("excerpt"):
                    ep = doc.add_paragraph(s["excerpt"][:300])
                    ep.runs[0].font.size = Pt(9) if ep.runs else None
                    ep.runs[0].font.color.rgb = RGBColor(100, 100, 100) if ep.runs else None
                if s.get("institution"):
                    doc.add_paragraph(f"机构: {s['institution']}")
                doc.add_paragraph()  # spacer

        # Disclaimer
        doc.add_paragraph()
        from app.policy.jurisdiction import get_disclaimer
        disclaimer_text = get_disclaimer(run.jurisdiction or "CN")
        footer = doc.add_paragraph()
        fr = footer.add_run(disclaimer_text)
        fr.italic = True
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor(170, 170, 170)
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

        filename = f"assistant_{run.run_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = output_dir / filename
        doc.save(output_path)
        return str(output_path)

    async def export_assistant_report_pdf(
        self,
        run,
        sources: list,
        mode_label: str = "法律分析",
        output_dir: Path = None,
    ) -> str:
        """Export assistant report as PDF (with WeasyPrint fallback to DOCX)."""
        if output_dir is None:
            output_dir = settings.storage_dir / "reports"
        output_dir.mkdir(parents=True, exist_ok=True)

        from app.policy.jurisdiction import get_disclaimer
        disclaimer_text = get_disclaimer(run.jurisdiction or "CN")

        sources_html = ""
        for s in sources:
            sid = s.get("source_id", "")
            trusted_label = "官方" if s.get("trusted") else "参考"
            color = "#28a745" if s.get("trusted") else "#e67e22"
            sources_html += f"""
            <div class="source">
                <span style="color:{color};font-weight:bold;">[{sid}] 【{trusted_label}】</span>
                <strong>{s.get('title', '')}</strong>
                <br/><small>{s.get('excerpt', '')[:200]}</small>
                {'<br/><a href="' + s['url'] + '">' + s['url'][:60] + '...</a>' if s.get('url') else ''}
            </div>
            """

        html = f"""
        <!DOCTYPE html><html><head><meta charset="utf-8">
        <title>法律助手报告</title>
        <style>
            body {{ font-family: "Microsoft YaHei", sans-serif; padding: 20px; }}
            h1 {{ text-align: center; color: #333; }}
            .info-table {{ width: 100%; border-collapse: collapse; margin: 20px 0; }}
            .info-table td {{ border: 1px solid #ddd; padding: 8px; }}
            .info-table td:first-child {{ background: #f5f5f5; width: 120px; }}
            .source {{ margin: 10px 0; padding: 10px; border: 1px solid #ddd; border-radius: 5px; }}
            .footer {{ text-align: center; color: #aaa; font-style: italic; margin-top: 30px; font-size: 10px; }}
        </style></head><body>
            <h1>法律助手 · {mode_label}报告</h1>
            <h2>基本信息</h2>
            <table class="info-table">
                <tr><td>报告类型</td><td>{mode_label}</td></tr>
                <tr><td>运行编号</td><td>{run.run_id if hasattr(run, 'run_id') else run.id}</td></tr>
                <tr><td>法域</td><td>{run.jurisdiction or 'CN'}</td></tr>
                <tr><td>生成时间</td><td>{run.created_at.strftime('%Y-%m-%d %H:%M:%S') if run.created_at else ''}</td></tr>
            </table>
            <h2>来源附录</h2>
            {sources_html}
            <p class="footer">{disclaimer_text}</p>
        </body></html>
        """

        try:
            from weasyprint import HTML
            filename = f"assistant_{run.run_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            output_path = output_dir / filename
            HTML(string=html).write_pdf(output_path)
            return str(output_path)
        except ImportError:
            logger.warning("WeasyPrint not available, falling back to DOCX")
            return await self.export_assistant_report_docx(run, sources, mode_label, output_dir)

    async def annotate_word_with_risks(
        self,
        file_content: bytes,
        risk_items: List[Dict[str, Any]],
        summary: str = "",
        output_dir: Path = None,
    ) -> str:
        """
        在用户上传的原始 Word 文档中添加 Word 原生审阅批注（Comment）。
        通过直接操作 DOCX ZIP 包来注入 word/comments.xml。
        """
        from docx import Document
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import RGBColor, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from lxml import etree
        import re
        import zipfile
        import shutil

        if output_dir is None:
            output_dir = settings.storage_dir / "redline"
        output_dir.mkdir(parents=True, exist_ok=True)

        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

        def normalize(text: str) -> str:
            return re.sub(r"\s+", "", text.strip())

        def longest_common_substring_len(s1: str, s2: str) -> int:
            """计算两个字符串的最长公共子串长度（DP）。"""
            if not s1 or not s2:
                return 0
            # 限制长度避免内存爆炸
            s1 = s1[:500]
            s2 = s2[:500]
            prev = [0] * (len(s2) + 1)
            best = 0
            for i in range(1, len(s1) + 1):
                curr = [0] * (len(s2) + 1)
                for j in range(1, len(s2) + 1):
                    if s1[i - 1] == s2[j - 1]:
                        curr[j] = prev[j - 1] + 1
                        if curr[j] > best:
                            best = curr[j]
                prev = curr
            return best

        def match_score(clause_norm: str, para_norm: str) -> float:
            """计算 clause 与 paragraph 的匹配得分 (0~1)。"""
            if not clause_norm or not para_norm:
                return 0.0
            # 完全包含：最高分
            if clause_norm in para_norm:
                return 1.0
            if para_norm in clause_norm:
                return 0.95
            # 最长公共子串占 clause 长度的比例
            lcs = longest_common_substring_len(clause_norm, para_norm)
            ratio = lcs / max(len(clause_norm), 1)
            return ratio

        # ── Step 1: Load doc, match risks to paragraphs ──
        doc = Document(io.BytesIO(file_content))

        severity_highlight = {"high": "red", "medium": "yellow", "low": "cyan"}

        # 预处理所有段落的 normalized text
        para_norms = []
        for para in doc.paragraphs:
            para_norms.append(normalize(para.text))

        # Match risks → paragraphs (用相似度评分找最佳匹配)
        matches = []  # [(para_index, risk, comment_id)]
        matched_para_indices = set()  # 避免同一段落被多条 risk 重复匹配
        comment_id = 0

        for risk in risk_items:
            clause_text = risk.get("clause_text", "")
            if not clause_text or len(clause_text) < 4:
                continue
            clause_norm = normalize(clause_text)
            if len(clause_norm) < 4:
                continue

            best_idx = -1
            best_score = 0.3  # 最低阈值：至少30%匹配度

            for idx, pn in enumerate(para_norms):
                if not pn or len(pn) < 4:
                    continue
                if idx in matched_para_indices:
                    continue  # 已被其他 risk 匹配，跳过

                score = match_score(clause_norm, pn)
                if score > best_score:
                    best_score = score
                    best_idx = idx

            if best_idx >= 0:
                matches.append((best_idx, risk, comment_id))
                matched_para_indices.add(best_idx)
                comment_id += 1

        # If nothing matched, attach all risks to first non-empty paragraph
        if not matches and risk_items and doc.paragraphs:
            first_idx = 0
            for i, pn in enumerate(para_norms):
                if pn and len(pn) > 4:
                    first_idx = i
                    break
            for risk in risk_items:
                matches.append((first_idx, risk, comment_id))
                comment_id += 1

        # ── Step 2: Build comments.xml content ──
        R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
        comments_root = etree.Element(
            f"{{{W_NS}}}comments",
            nsmap={"w": W_NS, "r": R_NS},
        )

        for _, risk, cid in matches:
            suggestion = risk.get("suggestion", "")
            name = risk.get("name", "")

            comment_elem = etree.SubElement(comments_root, f"{{{W_NS}}}comment")
            comment_elem.set(f"{{{W_NS}}}id", str(cid))
            comment_elem.set(f"{{{W_NS}}}author", "合同哨兵")
            comment_elem.set(f"{{{W_NS}}}initials", "CS")
            comment_elem.set(f"{{{W_NS}}}date", datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ"))

            # 批注直接写修改后的文本
            comment_text = suggestion if suggestion else name

            p1 = etree.SubElement(comment_elem, f"{{{W_NS}}}p")
            r1 = etree.SubElement(p1, f"{{{W_NS}}}r")
            t1 = etree.SubElement(r1, f"{{{W_NS}}}t")
            t1.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
            t1.text = comment_text

        comments_xml = etree.tostring(
            comments_root, xml_declaration=True, encoding="UTF-8", standalone=True
        )

        # ── Step 3: Apply redline edits (Track Changes style) + comments ──
        author = "合同哨兵"
        rev_date = datetime.now().strftime("%Y-%m-%dT%H:%M:%SZ")

        for para_idx, risk, cid in matches:
            para = doc.paragraphs[para_idx]
            suggestion = risk.get("suggestion", "")
            severity = risk.get("severity", "medium")

            # --- Comment: attach comment to paragraph ---
            range_start = OxmlElement("w:commentRangeStart")
            range_start.set(qn("w:id"), str(cid))
            para._p.insert(0, range_start)

            range_end = OxmlElement("w:commentRangeEnd")
            range_end.set(qn("w:id"), str(cid))
            para._p.append(range_end)

            ref_run = OxmlElement("w:r")
            ref_rpr = OxmlElement("w:rPr")
            ref_style = OxmlElement("w:rStyle")
            ref_style.set(qn("w:val"), "CommentReference")
            ref_rpr.append(ref_style)
            ref_run.append(ref_rpr)
            ref = OxmlElement("w:commentReference")
            ref.set(qn("w:id"), str(cid))
            ref_run.append(ref)
            para._p.append(ref_run)

            # --- Redline: wrap original runs in <w:del> and append <w:ins> with new text ---
            if suggestion:
                # Collect all existing runs
                existing_runs = list(para._p.findall(qn("w:r")))
                # Also collect the original text for verification
                original_text = para.text

                # Create <w:del> wrapper with all original runs
                del_elem = OxmlElement("w:del")
                del_elem.set(qn("w:id"), str(1000 + cid))
                del_elem.set(qn("w:author"), author)
                del_elem.set(qn("w:date"), rev_date)

                for orig_run in existing_runs:
                    # Skip commentReference runs
                    if orig_run.find(qn("w:commentReference")) is not None:
                        continue

                    # Clone the run, change <w:t> to <w:delText>
                    del_run = OxmlElement("w:r")
                    # Copy rPr if exists
                    orig_rpr = orig_run.find(qn("w:rPr"))
                    if orig_rpr is not None:
                        import copy
                        del_run.append(copy.deepcopy(orig_rpr))
                    # Add strikethrough to rPr
                    drpr = del_run.find(qn("w:rPr"))
                    if drpr is None:
                        drpr = OxmlElement("w:rPr")
                        del_run.insert(0, drpr)
                    strike = OxmlElement("w:strike")
                    drpr.append(strike)

                    # Convert <w:t> to <w:delText>
                    for t_elem in orig_run.findall(qn("w:t")):
                        del_text = OxmlElement("w:delText")
                        del_text.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                        del_text.text = t_elem.text or ""
                        del_run.append(del_text)

                    del_elem.append(del_run)

                # Create <w:ins> with new text
                ins_elem = OxmlElement("w:ins")
                ins_elem.set(qn("w:id"), str(2000 + cid))
                ins_elem.set(qn("w:author"), author)
                ins_elem.set(qn("w:date"), rev_date)

                ins_run = OxmlElement("w:r")
                ins_rpr = OxmlElement("w:rPr")
                # Red color for inserted text
                ins_color = OxmlElement("w:color")
                ins_color.set(qn("w:val"), "FF0000")
                ins_rpr.append(ins_color)
                ins_run.append(ins_rpr)
                ins_t = OxmlElement("w:t")
                ins_t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
                ins_t.text = suggestion
                ins_run.append(ins_t)
                ins_elem.append(ins_run)

                # Remove original runs (but keep commentRange and pPr)
                for orig_run in existing_runs:
                    if orig_run.find(qn("w:commentReference")) is not None:
                        continue
                    para._p.remove(orig_run)

                # Insert del then ins before the commentRangeEnd
                range_end_elem = para._p.find(qn("w:commentRangeEnd"))
                if range_end_elem is not None:
                    range_end_elem.addprevious(del_elem)
                    range_end_elem.addprevious(ins_elem)
                else:
                    para._p.append(del_elem)
                    para._p.append(ins_elem)
            else:
                # No suggestion, just highlight
                hl_color = severity_highlight.get(severity, "yellow")
                for run in para.runs:
                    rpr = run._r.get_or_add_rPr()
                    for old_hl in rpr.findall(qn("w:highlight")):
                        rpr.remove(old_hl)
                    hl = OxmlElement("w:highlight")
                    hl.set(qn("w:val"), hl_color)
                    rpr.append(hl)

        # ── Step 4: Add detailed review summary on a NEW PAGE ──
        if risk_items or summary:
            severity_labels = {"high": "高风险", "medium": "中风险", "low": "低风险"}
            severity_colors_map = {
                "high": RGBColor(220, 53, 69),
                "medium": RGBColor(230, 126, 34),
                "low": RGBColor(52, 152, 219),
            }

            # 插入分页符 —— 审核总结新开一页
            page_break_para = doc.add_paragraph()
            pb_run = page_break_para.add_run()
            br_elem = OxmlElement("w:br")
            br_elem.set(qn("w:type"), "page")
            pb_run._r.append(br_elem)

            # 标题
            heading = doc.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            hr = heading.add_run("━━━━━ 合同哨兵 · 审核总结 ━━━━━")
            hr.bold = True
            hr.font.size = Pt(14)
            hr.font.color.rgb = RGBColor(0, 102, 204)

            # 统计
            hc = sum(1 for r in risk_items if r.get("severity") == "high")
            mc = sum(1 for r in risk_items if r.get("severity") == "medium")
            lc = sum(1 for r in risk_items if r.get("severity") == "low")
            stat_p = doc.add_paragraph()
            stat_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sr = stat_p.add_run(f"高风险 {hc}  |  中风险 {mc}  |  低风险 {lc}  |  共 {len(risk_items)} 项")
            sr.font.size = Pt(10)
            sr.bold = True
            sr.font.color.rgb = RGBColor(100, 100, 100)

            # 整体评估
            if summary:
                doc.add_paragraph()
                sum_title = doc.add_paragraph()
                str_ = sum_title.add_run("整体评估")
                str_.bold = True
                str_.font.size = Pt(12)
                str_.font.color.rgb = RGBColor(50, 50, 50)

                sum_body = doc.add_paragraph()
                sbr = sum_body.add_run(summary)
                sbr.font.size = Pt(10)
                sbr.font.color.rgb = RGBColor(60, 60, 60)

            # 逐条详细说明
            if risk_items:
                doc.add_paragraph()
                detail_title = doc.add_paragraph()
                dtr = detail_title.add_run("修改明细与理由")
                dtr.bold = True
                dtr.font.size = Pt(12)
                dtr.font.color.rgb = RGBColor(50, 50, 50)

                for i, risk in enumerate(risk_items, 1):
                    severity = risk.get("severity", "medium")
                    label = severity_labels.get(severity, "风险")
                    color = severity_colors_map.get(severity, RGBColor(100, 100, 100))
                    name = risk.get("name", "")
                    clause = risk.get("clause_text", "")
                    suggestion = risk.get("suggestion", "")
                    desc = risk.get("description", "")
                    legal = risk.get("legal_basis", "")

                    # 序号 + 风险等级 + 名称
                    item_title = doc.add_paragraph()
                    itr = item_title.add_run(f"{i}. 【{label}】{name}")
                    itr.bold = True
                    itr.font.size = Pt(10)
                    itr.font.color.rgb = color

                    # 原文
                    if clause:
                        cp = doc.add_paragraph()
                        cp.add_run("原文：").bold = True
                        cr = cp.add_run(clause)
                        cr.font.size = Pt(9)
                        cr.font.color.rgb = RGBColor(120, 120, 120)

                    # 修改为
                    if suggestion:
                        sp = doc.add_paragraph()
                        sp.add_run("修改为：").bold = True
                        sgr = sp.add_run(suggestion)
                        sgr.font.size = Pt(9)
                        sgr.font.color.rgb = RGBColor(39, 174, 96)

                    # 理由
                    if desc:
                        dp = doc.add_paragraph()
                        dp.add_run("理由：").bold = True
                        dr = dp.add_run(desc)
                        dr.font.size = Pt(9)
                        dr.font.color.rgb = RGBColor(80, 80, 80)

                    # 法律依据
                    if legal:
                        lp = doc.add_paragraph()
                        lp.add_run("法律依据：").bold = True
                        lr = lp.add_run(legal)
                        lr.font.size = Pt(9)
                        lr.font.color.rgb = RGBColor(100, 100, 160)

                    doc.add_paragraph()  # spacer between items

            # 免责声明
            ft = doc.add_paragraph()
            ft.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fr = ft.add_run("本批注由合同哨兵AI生成，仅供参考，不构成法律意见或律师执业服务。重要决策请咨询持证律师。")
            fr.italic = True
            fr.font.size = Pt(8)
            fr.font.color.rgb = RGBColor(170, 170, 170)

        # ── Step 5: Save doc, then inject comments.xml into the ZIP ──
        tmp_filename = f"_tmp_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        tmp_path = output_dir / tmp_filename
        doc.save(tmp_path)

        final_filename = f"redline_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        final_path = output_dir / final_filename

        # Repack the DOCX ZIP with comments.xml injected
        with zipfile.ZipFile(str(tmp_path), "r") as zin:
            with zipfile.ZipFile(str(final_path), "w", zipfile.ZIP_DEFLATED) as zout:
                for item in zin.infolist():
                    data = zin.read(item.filename)

                    if item.filename == "[Content_Types].xml":
                        # Add comments content type if missing
                        ct_tree = etree.fromstring(data)
                        ct_ns = "http://schemas.openxmlformats.org/package/2006/content-types"
                        found = False
                        for override in ct_tree.findall(f"{{{ct_ns}}}Override"):
                            if override.get("PartName") == "/word/comments.xml":
                                found = True
                                break
                        if not found:
                            override_elem = etree.SubElement(ct_tree, f"{{{ct_ns}}}Override")
                            override_elem.set("PartName", "/word/comments.xml")
                            override_elem.set("ContentType",
                                "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml")
                        data = etree.tostring(ct_tree, xml_declaration=True, encoding="UTF-8", standalone=True)

                    elif item.filename == "word/_rels/document.xml.rels":
                        # Add comments relationship if missing
                        rels_tree = etree.fromstring(data)
                        rels_ns = "http://schemas.openxmlformats.org/package/2006/relationships"
                        comments_rel_type = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
                        found = False
                        max_rid = 0
                        for rel in rels_tree.findall(f"{{{rels_ns}}}Relationship"):
                            rid_str = rel.get("Id", "rId0")
                            try:
                                rid_num = int(rid_str.replace("rId", ""))
                                if rid_num > max_rid:
                                    max_rid = rid_num
                            except ValueError:
                                pass
                            if rel.get("Type") == comments_rel_type:
                                found = True
                                break
                        if not found:
                            new_rel = etree.SubElement(rels_tree, f"{{{rels_ns}}}Relationship")
                            new_rel.set("Id", f"rId{max_rid + 1}")
                            new_rel.set("Type", comments_rel_type)
                            new_rel.set("Target", "comments.xml")
                        data = etree.tostring(rels_tree, xml_declaration=True, encoding="UTF-8", standalone=True)

                    zout.writestr(item, data)

                # Write comments.xml
                zout.writestr("word/comments.xml", comments_xml)

        # Clean up temp file
        try:
            tmp_path.unlink()
        except Exception:
            pass

        logger.info(f"Redline Word saved: {final_path} ({len(matches)} comments)")
        return str(final_path)
