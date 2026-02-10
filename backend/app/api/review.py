"""Contract review API routes - 真实 AI 流水线 + Provenance/Verification 集成."""
import asyncio
import hashlib
import json
import time
import traceback
from datetime import datetime, timezone
from typing import AsyncGenerator, Optional

from fastapi import APIRouter, Depends, HTTPException, status, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.database import get_db, async_session_maker
from app.models.user import User
from app.models.contract import Contract, ContractStatus, ContractType, Jurisdiction
from app.models.review import ReviewResult
from app.schemas.review import ReviewRequest, ReviewResponse, ReviewProgressEvent
from app.api.deps import get_current_user, get_current_user_optional
from app.providers import get_provider_registry, ChatMessage
from app.config import settings
from app.pipeline.context import PipelineContext
from app.pipeline.verification import verify_review_output, verify_redline_output, get_verification_decision
from loguru import logger

router = APIRouter()


async def real_review_stream(
    contract_text: str,
    contract_type: str,
    jurisdiction: str,
    party_role: str,
    power_dynamic: str,
    ctx: Optional[dict] = None,
) -> AsyncGenerator[str, None]:
    """
    真实 AI 审核流水线 - 流式返回每个 Agent 的进度和结果
    
    Pipeline (ValueCell 模式):
      Stage 1: DocIngest (文档解析) → 已完成(上传时)
      Stage 2: ClauseStruct (条款结构化) → DeepSeek
      Stage 3: RuleEngine (规则预筛) → 本地规则
      Stage 4: LLMReview (AI深度审核) → DeepSeek
      Stage 5: RedlineDraft (修改建议生成) → DeepSeek
    
    ctx: mutable dict to collect results for persistence after streaming.
    """
    registry = get_provider_registry()
    all_risks = []
    total_tokens = 0

    # ── PipelineContext for provenance & verification ──
    pctx = PipelineContext(
        feature="review",
        user_id=ctx.get("user_id") if ctx else None,
        jurisdiction=jurisdiction,
        input_text=contract_text[:5000],
    )
    
    try:
        # ============================================
        # Stage 1: 文档解析 (已完成)
        # ============================================
        pctx.add_event("doc_ingest", "completed", 10, f"文本长度: {len(contract_text)} 字符")
        yield _sse({
            "stage": "doc_ingest",
            "status": "completed",
            "progress": 10,
            "message": "文档解析完成",
            "agent": "文档解析",
            "detail": f"文本长度: {len(contract_text)} 字符",
            "run_id": pctx.run_id,
        })
        await asyncio.sleep(0.3)
        
        # ============================================
        # Stage 2: 条款结构化 (ClauseStructAgent)
        # ============================================
        yield _sse({
            "stage": "clause_struct",
            "status": "running",
            "progress": 15,
            "message": "正在提取合同结构...",
            "agent": "条款结构化",
        })
        
        # Stage 2 用 MiniMax（结构提取是简单任务，节省 DeepSeek 额度）
        struct_client = registry.get_chat_client("minimax", "abab6.5s-chat")
        
        struct_prompt = f"""提取以下合同的关键信息，返回JSON：
{{
  "parties": ["甲方名称", "乙方名称"],
  "effective_date": "生效日期",
  "expiry_date": "到期日期",
  "contract_amount": "合同金额",
  "key_clauses": ["付款条款", "违约条款", "保密条款", ...]
}}

合同文本（前3000字）：
{contract_text[:3000]}"""
        
        struct_response = await struct_client.chat(
            [ChatMessage(role="user", content=struct_prompt)],
            temperature=0.1, max_tokens=1000
        )
        
        # 解析结构
        clauses = _normalize_clauses(_parse_json(struct_response.content))
        parties = clauses.get("parties", ["甲方", "乙方"])
        total_tokens += struct_response.tokens_input + struct_response.tokens_output
        
        yield _sse({
            "stage": "clause_struct",
            "status": "completed",
            "progress": 30,
            "message": f"结构提取完成 - 识别到 {len(parties)} 方当事人",
            "agent": "条款结构化",
            "detail": json.dumps(clauses, ensure_ascii=False),
            "tokens": struct_response.tokens_input + struct_response.tokens_output,
        })
        await asyncio.sleep(0.3)
        
        # ============================================
        # Stage 3: 规则预筛 (RuleEngineAgent) 
        # ============================================
        yield _sse({
            "stage": "rule_engine",
            "status": "running",
            "progress": 35,
            "message": "正在进行规则预筛...",
            "agent": "规则预筛",
        })
        
        from app.agents.rule_engine import RuleEngineAgent
        rule_engine = RuleEngineAgent()
        rule_matches = rule_engine.check(contract_text, jurisdiction=jurisdiction, contract_type=contract_type)
        
        for match in rule_matches:
            risk = {
                "id": match.rule_id,
                "severity": match.severity,
                "name": match.name,
                "description": match.description,
                "clause_text": match.matched_text,
                "suggestion": match.suggestion,
                "source": "rule_engine",
                "legal_basis": None,
            }
            all_risks.append(risk)
            yield _sse({
                "stage": "rule_engine",
                "status": "found_risk",
                "progress": 40,
                "message": f"规则发现: {match.name}",
                "agent": "规则预筛",
                "risk_item": risk,
            })
            await asyncio.sleep(0.2)
        
        yield _sse({
            "stage": "rule_engine",
            "status": "completed",
            "progress": 45,
            "message": f"规则预筛完成 - 发现 {len(rule_matches)} 个问题",
            "agent": "规则预筛",
        })
        await asyncio.sleep(0.3)
        
        # ============================================
        # Stage 3.5: Agent Search (实时搜取最新法条)
        # ============================================
        rag_context = ""
        try:
            yield _sse({
                "stage": "agent_search",
                "status": "running",
                "progress": 47,
                "message": "正在实时搜取最新法规...",
                "agent": "法规检索",
            })

            from app.rag.agent_search import AgentSearch
            agent_search = AgentSearch()
            search_results = await agent_search.search_laws(
                contract_type=contract_type,
                jurisdiction=jurisdiction,
                key_clauses=clauses.get("key_clauses", []),
            )

            if search_results:
                rag_parts = []
                for r in search_results:
                    source = r.get("source", "法规")
                    title = r.get("title", "")
                    text = r.get("text", "")[:300]
                    url = r.get("url", "")
                    rag_parts.append(f"【{source}】{title}\n{text}\n来源: {url}")
                rag_context = "\n\n".join(rag_parts)

            # 如果实时搜取没结果，回退到本地 RAG
            if not rag_context:
                try:
                    from app.rag.retriever import RagRetriever
                    async with async_session_maker() as rag_db:
                        retriever = RagRetriever(rag_db)
                        local_results = await retriever.search(
                            f"{contract_type} {jurisdiction}",
                            top_k=3, jurisdiction=jurisdiction,
                        )
                        if local_results:
                            rag_parts = [f"【本地知识库】{r['text'][:300]}" for r in local_results]
                            rag_context = "\n\n".join(rag_parts)
                except Exception:
                    pass

            result_count = len(search_results) if search_results else 0
            if rag_context:
                yield _sse({
                    "stage": "agent_search",
                    "status": "completed",
                    "progress": 49,
                    "message": f"搜取到 {result_count} 条相关法规",
                    "agent": "法规检索",
                })
            else:
                yield _sse({
                    "stage": "agent_search",
                    "status": "completed",
                    "progress": 49,
                    "message": "未搜取到相关法规",
                    "agent": "法规检索",
                })
        except Exception as search_err:
            logger.warning(f"Agent Search failed in review: {search_err}")
            yield _sse({
                "stage": "agent_search",
                "status": "completed",
                "progress": 49,
                "message": "法规搜取跳过（网络异常）",
                "agent": "法规检索",
            })
        await asyncio.sleep(0.2)

        # ============================================
        # Stage 4: AI 深度审核 (LLMReviewAgent) - 流式
        # ============================================
        chat_client = registry.get_chat_client("deepseek", "deepseek-reasoner")

        yield _sse({
            "stage": "llm_review",
            "status": "running",
            "progress": 50,
            "message": "深度审核中...",
            "agent": "深度审核",
        })
        
        # 构建专业审核 prompt（含甲乙方/强弱分析 + RAG 法规上下文）
        role_desc = _get_role_description(party_role, power_dynamic)
        existing_findings = "\n".join([f"- {r['name']}: {r['description']}" for r in all_risks[:5]])
        
        rag_section = ""
        if rag_context:
            rag_section = f"""
## 相关法规参考（来自知识库检索）
{rag_context[:2000]}
"""

        review_prompt = f"""你是一位资深法务审核专家。请对以下合同进行深度审核。

## 审核身份
{role_desc}

## 合同信息
- 类型: {contract_type}
- 适用法域: {jurisdiction}
- 当事人: {', '.join(parties)}

## 规则引擎已发现的问题（请勿重复）
{existing_findings if existing_findings else "暂无"}
{rag_section}
## 审核重点
1. 对我方不利的条款
2. 权利义务不对等
3. 隐含风险和陷阱
4. 法律合规性问题
5. 模糊不清需要澄清的条款

## 输出要求
逐条分析，每发现一个风险就立即输出，格式：
===RISK_START===
severity: high/medium/low
name: 风险名称
description: 详细描述（含对我方的具体影响）
clause_text: 合同原文（精确引用）
suggestion: 修改建议（具体可执行的修改文本）
legal_basis: 法律依据
===RISK_END===

分析完所有风险后，输出总结：
===SUMMARY_START===
整体评估内容
===SUMMARY_END===

## 合同全文
{contract_text[:10000]}"""
        
        # 流式获取 AI 审核结果
        full_response = ""
        async for token in chat_client.chat_stream(
            [ChatMessage(role="user", content=review_prompt)],
            temperature=0.2, max_tokens=4096
        ):
            full_response += token
            
            # 实时输出 token
            yield _sse({
                "stage": "llm_review",
                "status": "streaming",
                "progress": 55 + min(25, len(full_response) // 100),
                "message": "AI 正在分析...",
                "agent": "深度审核",
                "token": token,
            })
            
            # 检查是否有完整的风险项
            while "===RISK_END===" in full_response:
                risk_block, full_response = full_response.split("===RISK_END===", 1)
                if "===RISK_START===" in risk_block:
                    risk_text = risk_block.split("===RISK_START===")[1]
                    risk = _parse_risk_block(risk_text, len(all_risks))
                    if risk:
                        all_risks.append(risk)
                        yield _sse({
                            "stage": "llm_review",
                            "status": "found_risk",
                            "progress": 60,
                            "message": f"AI 发现: {risk['name']}",
                            "agent": "深度审核",
                            "risk_item": risk,
                        })
        
        # 提取总结
        summary = "审核完成"
        if "===SUMMARY_START===" in full_response:
            summary_parts = full_response.split("===SUMMARY_START===")
            if len(summary_parts) > 1:
                summary = summary_parts[1].split("===SUMMARY_END===")[0].strip()
        
        yield _sse({
            "stage": "llm_review",
            "status": "completed",
            "progress": 80,
            "message": f"AI 深度审核完成 - 共发现 {len(all_risks)} 个风险",
            "agent": "深度审核",
        })
        await asyncio.sleep(0.3)
        
        # ============================================
        # Stage 5: 修改建议生成 (RedlineDraftAgent) - 流式
        # ============================================
        high_medium = [r for r in all_risks if r["severity"] in ("high", "medium")]
        
        if high_medium:
            # Stage 5 继续用 DeepSeek Reasoner
            yield _sse({
                "stage": "redline_draft",
                "status": "running",
                "progress": 85,
                "message": "正在生成修改建议...",
                "agent": "修改建议",
            })
            
            redline_prompt = f"""针对以下合同风险，生成具体的修改建议文本。

## 审核身份
{role_desc}

## 需要修改的风险条款
{"".join([f'''
### {i+1}. {r["name"]} ({r["severity"]})
原文: {r["clause_text"]}
问题: {r["description"]}
''' for i, r in enumerate(high_medium[:5])])}

## 原始合同文本
{contract_text[:5000]}

请对每个风险条款，给出具体的修改后文本。格式：
===REDLINE_START===
risk_name: 风险名称
original: 原文
modified: 修改后文本
reason: 修改理由
===REDLINE_END==="""
            
            redline_full = ""
            redlines = []
            async for token in chat_client.chat_stream(
                [ChatMessage(role="user", content=redline_prompt)],
                temperature=0.3, max_tokens=3000
            ):
                redline_full += token
                yield _sse({
                    "stage": "redline_draft",
                    "status": "streaming",
                    "progress": 88,
                    "message": "正在生成修改文本...",
                    "agent": "修改建议",
                    "token": token,
                })
                
                while "===REDLINE_END===" in redline_full:
                    block, redline_full = redline_full.split("===REDLINE_END===", 1)
                    if "===REDLINE_START===" in block:
                        rl_text = block.split("===REDLINE_START===")[1]
                        rl = _parse_redline_block(rl_text)
                        if rl:
                            redlines.append(rl)
                            yield _sse({
                                "stage": "redline_draft",
                                "status": "found_redline",
                                "progress": 90,
                                "message": f"生成修改: {rl.get('risk_name', '')}",
                                "agent": "修改建议",
                                "redline": rl,
                            })
            
            yield _sse({
                "stage": "redline_draft",
                "status": "completed",
                "progress": 95,
                "message": f"修改建议生成完成 - {len(redlines)} 处修改",
                "agent": "修改建议",
            })
        
        # ============================================
        # 完成 - 持久化结果
        # ============================================
        high_count = sum(1 for r in all_risks if r["severity"] == "high")
        medium_count = sum(1 for r in all_risks if r["severity"] == "medium")
        low_count = sum(1 for r in all_risks if r["severity"] == "low")

        # Save results to database if ctx provided
        review_id = None
        if ctx is not None:
            try:
                review_id = await _save_review_result(
                    ctx=ctx,
                    all_risks=all_risks,
                    clauses=clauses,
                    summary=summary,
                    high_count=high_count,
                    medium_count=medium_count,
                    low_count=low_count,
                    total_tokens=total_tokens,
                )
            except Exception as e:
                from loguru import logger
                logger.warning(f"Failed to persist review result: {e}")

        # ── Verification ──
        pctx.add_event("verification", "running", 96, "验证中")
        verify_review_output(
            text=summary + " ".join(r.get("description", "") for r in all_risks),
            risk_items=all_risks,
            sources=[{"source_id": s["source_id"], "trusted": s["trusted"]} for s in pctx.sources],
            ctx=pctx,
        )
        v_decision = get_verification_decision(pctx)
        pctx.add_event("verification", "completed", 98, f"验证结果: {v_decision}")

        pctx.result_summary = {
            "high": high_count, "medium": medium_count, "low": low_count,
            "total": len(all_risks), "verification": v_decision,
        }
        try:
            await pctx.persist()
        except Exception as pe:
            logger.warning(f"Pipeline persist failed: {pe}")

        yield _sse({
            "stage": "complete",
            "status": "completed",
            "progress": 100,
            "message": "审核完成",
            "summary": summary,
            "review_id": review_id,
            "run_id": pctx.run_id,
            "verification_decision": v_decision,
            "stats": {
                "high": high_count,
                "medium": medium_count,
                "low": low_count,
                "total": len(all_risks),
            },
            "all_risks": all_risks,
        })
        
    except Exception as e:
        pctx.status = "failed"
        pctx.add_event("error", "error", 0, str(e))
        try:
            await pctx.persist()
        except Exception:
            pass
        yield _sse({
            "stage": "error",
            "status": "error",
            "progress": 0,
            "message": f"审核失败: {str(e)}",
            "detail": traceback.format_exc(),
        })


def _sse(data: dict) -> str:
    """Format as SSE event."""
    return f"data: {json.dumps(data, ensure_ascii=False)}\n\n"


def _get_role_description(party_role: str, power_dynamic: str) -> str:
    """生成审核身份描述（甲乙方 + 强弱势）"""
    role_map = {
        "party_a": "甲方（发起方/采购方）",
        "party_b": "乙方（供应方/服务方）",
        "third_party": "第三方独立审核",
    }
    power_map = {
        "strong": "强势方（如大厂/大客户）—— 重点审核法律合规和公平性，适度保护对方权益以促成合作",
        "weak": "弱势方（如小供应商/个人）—— 重点争取有利条款，识别对方的霸王条款",
        "equal": "对等谈判 —— 平衡双方权益，确保条款公平合理",
    }
    
    role = role_map.get(party_role, role_map["party_b"])
    power = power_map.get(power_dynamic, power_map["equal"])
    
    return f"""你站在 **{role}** 的角度审核合同。
谈判地位: {power}
审核策略: 根据我方地位，调整风险容忍度和建议力度。"""


def _parse_json(text: str) -> dict:
    """从文本中提取 JSON"""
    try:
        if "```json" in text:
            text = text.split("```json")[1].split("```")[0]
        elif "```" in text:
            text = text.split("```")[1].split("```")[0]
        return json.loads(text)
    except Exception:
        return {}


def _normalize_clauses(clauses: dict) -> dict:
    """
    归一化 AI 返回的 clauses 结构，确保所有字段类型正确。
    无论 AI 返回什么格式，都统一成 {"parties": [str, ...], "key_clauses": [str, ...]}
    """
    result = {}

    # parties: 确保是 str 列表
    raw_parties = clauses.get("parties", [])
    if isinstance(raw_parties, list):
        result["parties"] = [
            str(p.get("name", p)) if isinstance(p, dict) else str(p)
            for p in raw_parties
        ]
    elif isinstance(raw_parties, str):
        result["parties"] = [raw_parties]
    else:
        result["parties"] = ["甲方", "乙方"]

    # key_clauses: 确保是 str 列表
    raw_clauses = clauses.get("key_clauses", [])
    if isinstance(raw_clauses, list):
        result["key_clauses"] = [
            str(c.get("name", c.get("title", c))) if isinstance(c, dict) else str(c)
            for c in raw_clauses
        ]
    elif isinstance(raw_clauses, str):
        result["key_clauses"] = [raw_clauses]
    else:
        result["key_clauses"] = []

    # 保留其他字段
    for k, v in clauses.items():
        if k not in result:
            result[k] = v

    return result


def _parse_risk_block(text: str, index: int) -> dict:
    """解析风险块"""
    try:
        lines = text.strip().split("\n")
        risk = {"id": f"llm_{index}", "source": "llm_review"}
        for line in lines:
            line = line.strip()
            if line.startswith("severity:"):
                risk["severity"] = line.split(":", 1)[1].strip()
            elif line.startswith("name:"):
                risk["name"] = line.split(":", 1)[1].strip()
            elif line.startswith("description:"):
                risk["description"] = line.split(":", 1)[1].strip()
            elif line.startswith("clause_text:"):
                risk["clause_text"] = line.split(":", 1)[1].strip()
            elif line.startswith("suggestion:"):
                risk["suggestion"] = line.split(":", 1)[1].strip()
            elif line.startswith("legal_basis:"):
                risk["legal_basis"] = line.split(":", 1)[1].strip()
        
        if risk.get("name") and risk.get("severity"):
            return risk
    except:
        pass
    return None


def _parse_redline_block(text: str) -> dict:
    """解析修改建议块"""
    try:
        lines = text.strip().split("\n")
        rl = {}
        for line in lines:
            line = line.strip()
            if line.startswith("risk_name:"):
                rl["risk_name"] = line.split(":", 1)[1].strip()
            elif line.startswith("original:"):
                rl["original"] = line.split(":", 1)[1].strip()
            elif line.startswith("modified:"):
                rl["modified"] = line.split(":", 1)[1].strip()
            elif line.startswith("reason:"):
                rl["reason"] = line.split(":", 1)[1].strip()
        if rl.get("original") and rl.get("modified"):
            return rl
    except:
        pass
    return None


async def _save_review_result(
    ctx: dict,
    all_risks: list,
    clauses: dict,
    summary: str,
    high_count: int,
    medium_count: int,
    low_count: int,
    total_tokens: int,
) -> Optional[int]:
    """Persist Contract + ReviewResult to database after streaming completes."""
    async with async_session_maker() as db:
        # Create Contract record
        user_id = ctx.get("user_id")
        filename = ctx.get("filename", "unknown")
        file_content = ctx.get("file_content", b"")
        contract_type = ctx.get("contract_type", "general")
        jurisdiction = ctx.get("jurisdiction", "CN")
        contract_text = ctx.get("contract_text", "")

        file_hash = hashlib.sha256(file_content).hexdigest() if file_content else "none"

        # Save file to storage
        file_path = "in-memory"
        if user_id and file_content:
            user_dir = settings.storage_dir / str(user_id)
            user_dir.mkdir(parents=True, exist_ok=True)
            file_path = str(user_dir / f"{file_hash}_{filename}")
            with open(file_path, "wb") as f:
                f.write(file_content)

        # Safely convert string to enum
        try:
            ct_enum = ContractType(contract_type)
        except (ValueError, KeyError):
            ct_enum = ContractType.GENERAL
        try:
            j_enum = Jurisdiction(jurisdiction)
        except (ValueError, KeyError):
            j_enum = Jurisdiction.CN

        contract = Contract(
            user_id=user_id or 0,
            filename=filename,
            file_hash=file_hash,
            file_path=file_path,
            file_size=len(file_content),
            mime_type=ctx.get("mime_type", "application/octet-stream"),
            contract_type=ct_enum,
            jurisdiction=j_enum,
            raw_text=contract_text[:50000],
            status=ContractStatus.REVIEWED,
        )
        db.add(contract)
        await db.flush()

        # Create ReviewResult record
        review = ReviewResult(
            contract_id=contract.id,
            risk_items=all_risks,
            clauses=clauses,
            summary=summary,
            high_risk_count=high_count,
            medium_risk_count=medium_count,
            low_risk_count=low_count,
            model_used="deepseek-reasoner",
            tokens_used=total_tokens,
        )
        db.add(review)
        await db.commit()
        await db.refresh(review)
        return review.id


@router.get("/export/{review_id}")
async def export_review_report(
    review_id: int,
    format: str = "docx",
):
    """Export review result as DOCX or PDF."""
    from pathlib import Path
    from fastapi.responses import FileResponse
    from app.database import async_session_maker
    from app.services.export_service import ExportService

    async with async_session_maker() as db:
        result = await db.execute(
            select(ReviewResult).where(ReviewResult.id == review_id)
        )
        review = result.scalar_one_or_none()
        if not review:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="审核结果不存在",
            )

        contract_result = await db.execute(
            select(Contract).where(Contract.id == review.contract_id)
        )
        contract = contract_result.scalar_one_or_none()
        if not contract:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="合同不存在",
            )

        export_service = ExportService()

        # 用合同原名 + 后缀命名
        from urllib.parse import quote
        original_name = Path(contract.filename).stem  # 去掉扩展名

        if format == "pdf":
            output_path = await export_service.export_review_pdf(contract, review)
            media_type = "application/pdf"
            download_name = f"{original_name}_审核报告.pdf"
        else:
            output_path = await export_service.export_review_docx(contract, review)
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            download_name = f"{original_name}_审核报告.docx"

        encoded_name = quote(download_name)
        return FileResponse(
            path=output_path,
            media_type=media_type,
            filename="report.docx",
            headers={
                "Content-Disposition": f"attachment; filename=report.docx; filename*=UTF-8''{encoded_name}",
            },
        )


@router.post("/upload-and-review")
async def upload_and_review(
    file: UploadFile = File(...),
    contract_type: str = Form(default="general"),
    jurisdiction: str = Form(default="CN"),
    party_role: str = Form(default="party_b"),
    power_dynamic: str = Form(default="weak"),
    current_user: Optional[User] = Depends(get_current_user_optional),
):
    """上传合同并直接开始审核（流式返回），结果持久化到数据库"""
    # 读取文件内容
    content = await file.read()
    
    # 解析文档文本
    contract_text = ""
    filename = file.filename or ""
    
    if filename.endswith(".pdf"):
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(stream=content, filetype="pdf")
            for page in doc:
                contract_text += page.get_text()
            doc.close()
        except Exception as e:
            contract_text = f"[PDF解析失败: {e}]"
    elif filename.endswith(".docx"):
        try:
            import io
            from docx import Document
            doc = Document(io.BytesIO(content))
            for para in doc.paragraphs:
                contract_text += para.text + "\n"
        except Exception as e:
            contract_text = f"[DOCX解析失败: {e}]"
    else:
        # 尝试作为纯文本
        try:
            contract_text = content.decode("utf-8")
        except Exception:
            contract_text = content.decode("gbk", errors="ignore")
    
    if not contract_text.strip():
        contract_text = "[文档内容为空或无法解析]"

    # Context for persistence
    ctx = {
        "user_id": current_user.id if current_user else None,
        "filename": filename,
        "file_content": content,
        "mime_type": file.content_type or "application/octet-stream",
        "contract_type": contract_type,
        "jurisdiction": jurisdiction,
        "contract_text": contract_text,
    }
    
    return StreamingResponse(
        real_review_stream(
            contract_text=contract_text,
            contract_type=contract_type,
            jurisdiction=jurisdiction,
            party_role=party_role,
            power_dynamic=power_dynamic,
            ctx=ctx,
        ),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


async def redline_review_stream(
    contract_text: str,
    file_content: bytes,
    filename: str,
    contract_type: str,
    jurisdiction: str,
    party_role: str,
    power_dynamic: str,
) -> AsyncGenerator[str, None]:
    """批阅模式流水线 - 流式返回进度，最终生成批注版 Word 下载链接。"""
    import base64
    registry = get_provider_registry()
    all_risks = []
    summary = ""

    # ── PipelineContext for redline ──
    pctx = PipelineContext(
        feature="redline",
        jurisdiction=jurisdiction,
        input_text=contract_text[:5000],
    )

    try:
        # Stage 1: 文档解析
        pctx.add_event("doc_ingest", "completed", 10, f"文档解析完成 ({len(contract_text)} 字符)")
        yield _sse({"stage": "doc_ingest", "status": "completed", "progress": 10,
                     "message": f"文档解析完成 ({len(contract_text)} 字符)", "agent": "文档解析",
                     "run_id": pctx.run_id})
        await asyncio.sleep(0.3)

        # Stage 2: 条款结构化
        yield _sse({"stage": "clause_struct", "status": "running", "progress": 15,
                     "message": "正在提取合同结构...", "agent": "条款结构化"})

        struct_client = registry.get_chat_client("minimax", "abab6.5s-chat")
        struct_prompt = f"""提取以下合同的关键信息，返回JSON：
{{
  "parties": ["甲方名称", "乙方名称"],
  "key_clauses": ["付款条款", "违约条款", "保密条款", ...]
}}

合同文本（前3000字）：
{contract_text[:3000]}"""
        struct_response = await struct_client.chat(
            [ChatMessage(role="user", content=struct_prompt)],
            temperature=0.1, max_tokens=1000
        )
        clauses = _normalize_clauses(_parse_json(struct_response.content))
        parties = clauses.get("parties", ["甲方", "乙方"])

        yield _sse({"stage": "clause_struct", "status": "completed", "progress": 30,
                     "message": f"结构提取完成 - 识别到 {len(parties)} 方当事人", "agent": "条款结构化"})
        await asyncio.sleep(0.3)

        # Stage 3: 规则预筛
        yield _sse({"stage": "rule_engine", "status": "running", "progress": 35,
                     "message": "正在进行规则预筛...", "agent": "规则预筛"})

        from app.agents.rule_engine import RuleEngineAgent
        rule_engine = RuleEngineAgent()
        rule_matches = rule_engine.check(contract_text, jurisdiction=jurisdiction, contract_type=contract_type)
        for match in rule_matches:
            risk = {
                "id": match.rule_id, "severity": match.severity, "name": match.name,
                "description": match.description, "clause_text": match.matched_text,
                "suggestion": match.suggestion, "source": "rule_engine", "legal_basis": None,
            }
            all_risks.append(risk)
            yield _sse({"stage": "rule_engine", "status": "found_risk", "progress": 40,
                         "message": f"规则发现: {match.name}", "agent": "规则预筛", "risk_item": risk})
            await asyncio.sleep(0.15)

        yield _sse({"stage": "rule_engine", "status": "completed", "progress": 45,
                     "message": f"规则预筛完成 - 发现 {len(rule_matches)} 个问题", "agent": "规则预筛"})
        await asyncio.sleep(0.3)

        # Stage 3.5: Agent Search（实时搜取最新法条）
        rag_context = ""
        try:
            yield _sse({"stage": "agent_search", "status": "running", "progress": 47,
                         "message": "正在实时搜取最新法规...", "agent": "法规检索"})

            from app.rag.agent_search import AgentSearch
            agent_search = AgentSearch()
            search_results = await agent_search.search_laws(
                contract_type=contract_type,
                jurisdiction=jurisdiction,
                key_clauses=clauses.get("key_clauses", []),
            )
            if search_results:
                rag_parts = []
                for r in search_results:
                    source = r.get("source", "法规")
                    title = r.get("title", "")
                    text = r.get("text", "")[:300]
                    rag_parts.append(f"【{source}】{title}\n{text}")
                rag_context = "\n\n".join(rag_parts)

            # 回退到本地 RAG
            if not rag_context:
                try:
                    from app.rag.retriever import RagRetriever
                    async with async_session_maker() as rag_db:
                        retriever = RagRetriever(rag_db)
                        local_results = await retriever.search(f"{contract_type} {jurisdiction}", top_k=3)
                        if local_results:
                            rag_context = "\n\n".join([f"【本地】{r['text'][:300]}" for r in local_results])
                except Exception:
                    pass

            count = len(search_results) if search_results else 0
            msg = f"搜取到 {count} 条相关法规" if rag_context else "未搜取到相关法规"
            yield _sse({"stage": "agent_search", "status": "completed", "progress": 49,
                         "message": msg, "agent": "法规检索"})
        except Exception as search_err:
            logger.warning(f"Agent Search failed in redline: {search_err}")
            yield _sse({"stage": "agent_search", "status": "completed", "progress": 49,
                         "message": "法规搜取跳过", "agent": "法规检索"})
        await asyncio.sleep(0.2)

        # Stage 4: 深度审核
        yield _sse({"stage": "llm_review", "status": "running", "progress": 50,
                     "message": "深度审核中...", "agent": "深度审核"})

        chat_client = registry.get_chat_client("deepseek", "deepseek-reasoner")
        role_desc = _get_role_description(party_role, power_dynamic)
        existing_findings = "\n".join([f"- {r['name']}: {r['description']}" for r in all_risks[:5]])

        rag_section = ""
        if rag_context:
            rag_section = f"""
## 相关法规参考
{rag_context[:2000]}
"""

        review_prompt = f"""你是一位资深法务总监，正在直接修改一份合同。你要像真正的法务一样，直接把有问题的条款改掉。

## 你的身份
{role_desc}

## 合同信息
- 类型: {contract_type}  |  法域: {jurisdiction}  |  当事人: {', '.join(parties)}

## 已发现问题（勿重复）
{existing_findings if existing_findings else "无"}
{rag_section}
## 工作方式
你要做的是**直接修改合同**，不是提建议。具体：
1. 找到有问题的条款
2. clause_text 必须**逐字复制**原文中的完整句子（不改一个字，用于精确定位）
3. suggestion 写出**修改后的完整替换文本**（直接替换原文的最终版本，不要写"建议"、"应"、"可以改为"等词）
4. description 和 legal_basis 只在审核总结中展示，不会出现在修改标记中

## 输出格式
===RISK_START===
severity: high/medium/low
name: 修改点名称（3-8字）
clause_text: 原文（逐字复制）
suggestion: 修改后文本（直接替换用）
description: 修改理由
legal_basis: 法条
===RISK_END===

最后输出总结：
===SUMMARY_START===
整体风险评价、谈判策略
===SUMMARY_END===

## 合同全文
{contract_text[:10000]}"""

        full_response = ""
        async for token in chat_client.chat_stream(
            [ChatMessage(role="user", content=review_prompt)],
            temperature=0.2, max_tokens=4096
        ):
            full_response += token
            yield _sse({"stage": "llm_review", "status": "streaming", "progress": 55 + min(25, len(full_response) // 100),
                         "message": "深度分析中...", "agent": "深度审核", "token": token})

            while "===RISK_END===" in full_response:
                risk_block, full_response = full_response.split("===RISK_END===", 1)
                if "===RISK_START===" in risk_block:
                    risk_text = risk_block.split("===RISK_START===")[1]
                    risk = _parse_risk_block(risk_text, len(all_risks))
                    if risk:
                        all_risks.append(risk)
                        yield _sse({"stage": "llm_review", "status": "found_risk", "progress": 65,
                                     "message": f"发现: {risk['name']}", "agent": "深度审核", "risk_item": risk})

        if "===SUMMARY_START===" in full_response:
            parts = full_response.split("===SUMMARY_START===")
            if len(parts) > 1:
                summary = parts[1].split("===SUMMARY_END===")[0].strip()

        yield _sse({"stage": "llm_review", "status": "completed", "progress": 85,
                     "message": f"深度审核完成 - 共发现 {len(all_risks)} 个风险", "agent": "深度审核"})
        await asyncio.sleep(0.3)

        # Stage 5: 生成批注文档
        yield _sse({"stage": "annotate", "status": "running", "progress": 90,
                     "message": "正在生成批注文档...", "agent": "文档批注"})

        download_filename = None
        try:
            from app.services.export_service import ExportService
            export_service = ExportService()
            output_path = await export_service.annotate_word_with_risks(
                file_content=file_content,
                risk_items=all_risks,
                summary=summary or "审核完成",
            )
            from pathlib import Path as _Path
            download_filename = _Path(output_path).name

            yield _sse({"stage": "annotate", "status": "completed", "progress": 95,
                         "message": "批注文档生成完成", "agent": "文档批注"})
        except Exception as ann_err:
            logger.error(f"Annotation failed: {ann_err}")
            yield _sse({"stage": "annotate", "status": "error", "progress": 95,
                         "message": f"批注文档生成失败: {str(ann_err)}", "agent": "文档批注"})

        # Stage 6: 完成 - 始终发送 complete 事件
        high_count = sum(1 for r in all_risks if r.get("severity") == "high")
        medium_count = sum(1 for r in all_risks if r.get("severity") == "medium")
        low_count = sum(1 for r in all_risks if r.get("severity") == "low")

        import os
        original_stem = os.path.splitext(filename)[0]
        user_download_name = f"{original_stem}_批注版.docx"

        # ── Verification (ClauseLocateVerify) ──
        pctx.add_event("verification", "running", 96, "验证批注定位")
        doc_paragraphs = [p.strip() for p in contract_text.split("\n") if p.strip()]
        verify_redline_output(
            risk_items=all_risks,
            doc_paragraphs=doc_paragraphs,
            sources=[{"source_id": s["source_id"], "trusted": s["trusted"]} for s in pctx.sources],
            ctx=pctx,
        )
        v_decision = get_verification_decision(pctx)
        pctx.add_event("verification", "completed", 98, f"验证结果: {v_decision}")

        pctx.result_summary = {
            "high": high_count, "medium": medium_count, "low": low_count,
            "total": len(all_risks), "verification": v_decision,
        }
        try:
            await pctx.persist()
        except Exception as pe:
            logger.warning(f"Pipeline persist failed: {pe}")

        complete_data = {
            "stage": "complete", "status": "completed", "progress": 100,
            "message": "批阅完成",
            "summary": summary,
            "run_id": pctx.run_id,
            "verification_decision": v_decision,
            "stats": {"high": high_count, "medium": medium_count, "low": low_count, "total": len(all_risks)},
            "all_risks": all_risks,
        }
        if download_filename:
            complete_data["download_url"] = f"/api/review/redline/download/{download_filename}"
            complete_data["download_name"] = user_download_name

        yield _sse(complete_data)

    except Exception as e:
        pctx.status = "failed"
        pctx.add_event("error", "error", 0, str(e))
        try:
            await pctx.persist()
        except Exception:
            pass
        high_count = sum(1 for r in all_risks if r.get("severity") == "high")
        medium_count = sum(1 for r in all_risks if r.get("severity") == "medium")
        low_count = sum(1 for r in all_risks if r.get("severity") == "low")
        yield _sse({
            "stage": "error", "status": "error", "progress": 0,
            "message": f"批阅失败: {str(e)}",
            "run_id": pctx.run_id,
            "stats": {"high": high_count, "medium": medium_count, "low": low_count, "total": len(all_risks)},
            "all_risks": all_risks,
            "detail": traceback.format_exc(),
        })


@router.get("/redline/download/{filename}")
async def download_redline(filename: str, name: Optional[str] = None):
    """Download a generated redline Word document."""
    from pathlib import Path
    from fastapi.responses import FileResponse
    from urllib.parse import quote

    file_path = settings.storage_dir / "redline" / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在或已过期")

    download_name = name or filename
    encoded = quote(download_name)
    return FileResponse(
        path=str(file_path),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename="download.docx",
        headers={"Content-Disposition": f"attachment; filename=download.docx; filename*=UTF-8''{encoded}"},
    )


@router.post("/redline")
async def upload_and_redline(
    file: UploadFile = File(...),
    contract_type: str = Form(default="general"),
    jurisdiction: str = Form(default="CN"),
    party_role: str = Form(default="party_b"),
    power_dynamic: str = Form(default="weak"),
    current_user: Optional[User] = Depends(get_current_user_optional),
):
    """上传 Word → 流式 AI 批阅 → 生成批注版 Word（SSE 流式进度）"""
    filename = file.filename or "contract.docx"
    if not filename.endswith(".docx"):
        raise HTTPException(status_code=400, detail="批注功能仅支持 .docx 格式")

    content = await file.read()

    try:
        import io as _io
        from docx import Document as DocxDocument
        doc = DocxDocument(_io.BytesIO(content))
        contract_text = "\n".join(para.text for para in doc.paragraphs)
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Word 解析失败: {e}")

    if not contract_text.strip():
        raise HTTPException(status_code=400, detail="文档内容为空")

    return StreamingResponse(
        redline_review_stream(
            contract_text=contract_text,
            file_content=content,
            filename=filename,
            contract_type=contract_type,
            jurisdiction=jurisdiction,
            party_role=party_role,
            power_dynamic=power_dynamic,
        ),
        media_type="text/event-stream",
        headers={"Cache-Control": "no-cache", "Connection": "keep-alive", "X-Accel-Buffering": "no"},
    )
