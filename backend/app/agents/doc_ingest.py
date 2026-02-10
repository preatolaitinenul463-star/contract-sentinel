"""Document ingestion agent - parses PDF/DOCX/images."""
import hashlib
from dataclasses import dataclass, field
from pathlib import Path
from typing import List, Optional, Dict, Any

from loguru import logger


@dataclass
class ClauseLocation:
    """Location of text within document."""
    page: Optional[int] = None
    paragraph: Optional[int] = None
    start_char: Optional[int] = None
    end_char: Optional[int] = None
    
    def to_dict(self) -> Dict[str, Any]:
        return {
            "page": self.page,
            "paragraph": self.paragraph,
            "start": self.start_char,
            "end": self.end_char,
        }


@dataclass
class TextBlock:
    """A block of text with location info."""
    text: str
    location: ClauseLocation
    block_type: str = "paragraph"  # paragraph, heading, list_item, table_cell


@dataclass
class ParsedDocument:
    """Result of document parsing."""
    filename: str
    mime_type: str
    page_count: int
    raw_text: str
    blocks: List[TextBlock] = field(default_factory=list)
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def get_text_with_locations(self) -> List[Dict[str, Any]]:
        """Get text blocks with location info."""
        return [
            {
                "text": block.text,
                "type": block.block_type,
                "location": block.location.to_dict(),
            }
            for block in self.blocks
        ]


class DocIngestAgent:
    """Agent for parsing various document formats."""
    
    def __init__(self):
        self._ocr_engine = None
    
    async def parse(self, file_path: str, mime_type: str) -> ParsedDocument:
        """Parse a document file."""
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        logger.info(f"Parsing document: {path.name} ({mime_type})")
        
        if mime_type == "application/pdf":
            return await self._parse_pdf(path)
        elif mime_type in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        ):
            return await self._parse_docx(path)
        elif mime_type.startswith("image/"):
            return await self._parse_image(path)
        else:
            raise ValueError(f"Unsupported file type: {mime_type}")
    
    async def _parse_pdf(self, path: Path) -> ParsedDocument:
        """Parse PDF document using PyMuPDF."""
        import fitz  # PyMuPDF
        
        doc = fitz.open(path)
        blocks: List[TextBlock] = []
        full_text_parts: List[str] = []
        
        try:
            for page_num, page in enumerate(doc):
                # Extract text blocks with position info
                text_dict = page.get_text("dict")
                
                for block_idx, block in enumerate(text_dict.get("blocks", [])):
                    if block.get("type") == 0:  # Text block
                        block_text = ""
                        for line in block.get("lines", []):
                            for span in line.get("spans", []):
                                block_text += span.get("text", "")
                            block_text += "\n"
                        
                        block_text = block_text.strip()
                        if block_text:
                            text_block = TextBlock(
                                text=block_text,
                                location=ClauseLocation(
                                    page=page_num + 1,
                                    paragraph=block_idx + 1,
                                    start_char=len("\n\n".join(full_text_parts)),
                                    end_char=len("\n\n".join(full_text_parts)) + len(block_text),
                                ),
                                block_type="paragraph",
                            )
                            blocks.append(text_block)
                            full_text_parts.append(block_text)
            
            raw_text = "\n\n".join(full_text_parts)
            
            return ParsedDocument(
                filename=path.name,
                mime_type="application/pdf",
                page_count=len(doc),
                raw_text=raw_text,
                blocks=blocks,
                metadata={
                    "title": doc.metadata.get("title", ""),
                    "author": doc.metadata.get("author", ""),
                    "creation_date": doc.metadata.get("creationDate", ""),
                },
            )
        finally:
            doc.close()
    
    async def _parse_docx(self, path: Path) -> ParsedDocument:
        """Parse DOCX document using python-docx."""
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError
        
        try:
            doc = Document(path)
        except PackageNotFoundError:
            raise ValueError("Invalid DOCX file")
        
        blocks: List[TextBlock] = []
        full_text_parts: List[str] = []
        
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if text:
                # Determine block type based on style
                block_type = "paragraph"
                style_name = paragraph.style.name if paragraph.style else ""
                if style_name.startswith("Heading"):
                    block_type = "heading"
                elif style_name.startswith("List"):
                    block_type = "list_item"
                
                text_block = TextBlock(
                    text=text,
                    location=ClauseLocation(
                        page=None,  # DOCX doesn't have inherent page numbers
                        paragraph=para_idx + 1,
                        start_char=len("\n\n".join(full_text_parts)),
                        end_char=len("\n\n".join(full_text_parts)) + len(text),
                    ),
                    block_type=block_type,
                )
                blocks.append(text_block)
                full_text_parts.append(text)
        
        # Also extract tables
        for table_idx, table in enumerate(doc.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    if text:
                        text_block = TextBlock(
                            text=text,
                            location=ClauseLocation(
                                paragraph=len(blocks) + 1,
                            ),
                            block_type="table_cell",
                        )
                        blocks.append(text_block)
                        full_text_parts.append(text)
        
        raw_text = "\n\n".join(full_text_parts)
        
        # Get core properties
        metadata = {}
        try:
            core_props = doc.core_properties
            metadata = {
                "title": core_props.title or "",
                "author": core_props.author or "",
                "created": str(core_props.created) if core_props.created else "",
            }
        except Exception:
            pass
        
        return ParsedDocument(
            filename=path.name,
            mime_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            page_count=1,  # DOCX page count requires rendering
            raw_text=raw_text,
            blocks=blocks,
            metadata=metadata,
        )
    
    async def _parse_image(self, path: Path) -> ParsedDocument:
        """Parse image using OCR (RapidOCR)."""
        # Lazy load OCR engine
        if self._ocr_engine is None:
            try:
                from rapidocr_onnxruntime import RapidOCR
                self._ocr_engine = RapidOCR()
            except ImportError:
                raise ImportError("RapidOCR not installed. Run: pip install rapidocr-onnxruntime")
        
        # Run OCR
        result, _ = self._ocr_engine(str(path))
        
        blocks: List[TextBlock] = []
        full_text_parts: List[str] = []
        
        if result:
            for idx, (box, text, confidence) in enumerate(result):
                if text and confidence > 0.5:  # Filter low confidence
                    text_block = TextBlock(
                        text=text,
                        location=ClauseLocation(
                            page=1,
                            paragraph=idx + 1,
                            start_char=len("\n".join(full_text_parts)),
                            end_char=len("\n".join(full_text_parts)) + len(text),
                        ),
                        block_type="paragraph",
                    )
                    blocks.append(text_block)
                    full_text_parts.append(text)
        
        raw_text = "\n".join(full_text_parts)
        
        return ParsedDocument(
            filename=path.name,
            mime_type="image/ocr",
            page_count=1,
            raw_text=raw_text,
            blocks=blocks,
            metadata={"ocr": True},
        )
